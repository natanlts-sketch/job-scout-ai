from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from src.core.text import safe_filename


def find_matching_keywords(
    cv_keywords: list[str],
    job_keywords: list[str],
    limit: int = 12,
) -> list[str]:
    cv_set = {keyword.lower() for keyword in cv_keywords}
    matches = []
    for keyword in job_keywords:
        normalized = keyword.lower()
        if normalized in cv_set and normalized not in matches:
            matches.append(normalized)
    return matches[:limit]


def _rewrite_summary(existing: str, job_title: str, company: str, skills: list[str]) -> str:
    skill_bit = ", ".join(s.title() for s in skills[:6]) if skills else "data analysis"
    base = existing.strip() if existing else "Junior data professional"
    # Keep truthful: do not invent employers or years
    return (
        f"{base} Seeking a {job_title} role at {company}. "
        f"Relevant strengths include {skill_bit}."
    )[:500]


def create_tailored_cv(
    master_cv_path: Path,
    output_directory: Path,
    job_title: str,
    company: str,
    matching_keywords: list[str],
    *,
    rewrite_summary: bool = True,
) -> Path:
    """
    Create a job-specific copy of the master CV.

    Preserves original formatting; prepends truthful relevant skills and
    optionally adjusts the first short paragraph as a summary hint.
    Never invents skills beyond matching_keywords (already CV-validated).
    """
    document = Document(master_cv_path)

    if matching_keywords:
        # Reorder: put matching skills first in a dedicated line
        skills_text = "Relevant Skills (for this role): " + ", ".join(
            keyword.title() for keyword in matching_keywords
        )
        paragraph = document.paragraphs[0].insert_paragraph_before(skills_text)
        paragraph.style = document.styles["Normal"]

        if rewrite_summary and len(document.paragraphs) > 2:
            # Soft rewrite of a short early paragraph if it looks like a summary
            for idx, para in enumerate(document.paragraphs[1:6], start=1):
                text = para.text.strip()
                if 40 < len(text) < 400 and not text.lower().startswith("relevant skills"):
                    para.text = _rewrite_summary(text, job_title, company, matching_keywords)
                    break

    output_directory.mkdir(parents=True, exist_ok=True)
    filename = (
        f"{safe_filename(company)}_"
        f"{safe_filename(job_title)}_CV.docx"
    )
    output_path = output_directory / filename
    document.save(output_path)
    return output_path


def compare_cv_text(original_text: str, tailored_path: Path) -> dict:
    from src.cv.parser import CVParser

    tailored_text = CVParser(tailored_path).get_text()
    return {
        "original_chars": len(original_text),
        "tailored_chars": len(tailored_text),
        "original_preview": original_text[:800],
        "tailored_preview": tailored_text[:800],
    }


def export_pdf_from_docx(docx_path: Path, pdf_path: Path | None = None) -> Path:
    """Create a simple PDF text export from DOCX (layout-light, truthful content)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    from src.cv.parser import CVParser

    pdf_path = pdf_path or docx_path.with_suffix(".pdf")
    text = CVParser(docx_path).get_text()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    y = height - 40
    for line in text.splitlines():
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line[:110])
        y -= 14
    c.save()
    return pdf_path
