from __future__ import annotations

import re
from pathlib import Path

from src.core.logging_setup import get_logger

logger = get_logger("jobscout.cv.parser")


class CVParser:
    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.suffix = self.file_path.suffix.lower()

    def get_text(self) -> str:
        if self.suffix == ".docx":
            return self._from_docx()
        if self.suffix == ".pdf":
            return self._from_pdf()
        raise ValueError(f"Unsupported CV format: {self.suffix}")

    def _from_docx(self) -> str:
        from docx import Document

        document = Document(self.file_path)
        text = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())
        # tables
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text.append(" | ".join(cells))
        return "\n".join(text)

    def _from_pdf(self) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(self.file_path))
        parts = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted.strip():
                parts.append(extracted.strip())
        return "\n".join(parts)


SECTION_PATTERNS = {
    "experience": re.compile(
        r"(professional experience|work experience|experience|ניסיון תעסוקתי|ניסיון)",
        re.I,
    ),
    "education": re.compile(r"(education|academic|השכלה|לימודים)", re.I),
    "skills": re.compile(r"(skills|technical skills|כישורים|מיומנויות)", re.I),
    "languages": re.compile(r"(languages|שפות)", re.I),
    "summary": re.compile(r"(summary|profile|objective|תקציר|פרופיל)", re.I),
}


def extract_cv_sections(text: str) -> dict[str, str]:
    lines = text.splitlines()
    sections: dict[str, list[str]] = {k: [] for k in SECTION_PATTERNS}
    current = None
    for line in lines:
        matched_header = None
        for name, pattern in SECTION_PATTERNS.items():
            if pattern.search(line) and len(line) < 80:
                matched_header = name
                break
        if matched_header:
            current = matched_header
            continue
        if current:
            sections[current].append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items() if v}
