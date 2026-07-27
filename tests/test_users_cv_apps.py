from pathlib import Path

from src.auth import create_user, get_preferences, save_preferences, authenticate, delete_user_and_data
from src.core import db as dbmod
from src.core.db import initialize_database, mark_new_jobs_legacy
from src.core.models import Job
from src.cv.upload import add_user_skill, get_user_skills, store_cv_upload, validate_upload
from src.applications import create_application_package, list_applications
import pytest


@pytest.fixture()
def tmp_db(tmp_path, monkeypatch):
    db_path = tmp_path / "test.db"
    monkeypatch.setattr(dbmod, "get_db_path", lambda: db_path)
    # Also patch get_path used by uploads/applications via config — set env-less overrides
    from src.core import config as configmod

    original = configmod.get_path

    def fake_get_path(key: str):
        mapping = {
            "database": db_path,
            "uploads": tmp_path / "uploads",
            "applications": tmp_path / "applications",
            "reports": tmp_path / "reports",
            "tailored_cvs": tmp_path / "tailored",
            "backups": tmp_path / "backups",
            "master_cv": tmp_path / "master.docx",
        }
        if key in mapping:
            return mapping[key]
        return original(key)

    monkeypatch.setattr(configmod, "get_path", fake_get_path)
    # applications and upload import get_path at call time from config — patch those too
    import src.cv.upload as uploadmod
    import src.applications as appsmod
    import src.security as securomod

    monkeypatch.setattr(uploadmod, "get_path", fake_get_path)
    monkeypatch.setattr(appsmod, "get_path", fake_get_path)
    monkeypatch.setattr(securomod, "get_path", fake_get_path)
    initialize_database(db_path)
    return db_path


def test_user_separation(tmp_db):
    u1 = create_user("a@example.com", "password123", "A")
    u2 = create_user("b@example.com", "password123", "B")
    add_user_skill(u1["id"], "sql")
    add_user_skill(u2["id"], "python")
    assert get_user_skills(u1["id"]) == ["sql"]
    assert get_user_skills(u2["id"]) == ["python"]
    assert authenticate("a@example.com", "password123")
    assert authenticate("a@example.com", "wrong") is None


def test_preferences_roundtrip(tmp_db):
    user = create_user("prefs@example.com", "password123")
    save_preferences(
        user["id"],
        {
            "preferred_titles": ["Data Analyst"],
            "preferred_locations": ["Tel Aviv"],
            "remote_preference": "hybrid",
            "minimum_ats_score": 40,
            "excluded_companies": ["BadCorp"],
            "ai_consent": True,
            "email_notifications": False,
            "search_frequency_hours": 12,
        },
    )
    prefs = get_preferences(user["id"])
    assert prefs["preferred_titles"] == ["Data Analyst"]
    assert prefs["remote_preference"] == "hybrid"
    assert prefs["ai_consent"] in (1, True)
    assert prefs["email_notifications"] in (0, False)


def test_duplicate_job_detection(tmp_db):
    conn = initialize_database(tmp_db)
    result = mark_new_jobs_legacy(conn, ["abc"], "2026-07-20T00:00:00+00:00")
    assert result["abc"] is True
    result2 = mark_new_jobs_legacy(conn, ["abc"], "2026-07-20T00:00:00+00:00")
    assert result2["abc"] is False


def test_validate_upload_rejects_bad_type():
    with pytest.raises(ValueError):
        validate_upload("virus.exe", 100)


def test_cv_upload_docx(tmp_db, tmp_path):
    from docx import Document

    user = create_user("cv@example.com", "password123")
    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Natan Mamedov")
    doc.add_paragraph("Skills: SQL Python Tableau Excel Power BI")
    doc.save(docx_path)
    result = store_cv_upload(user["id"], "sample.docx", docx_path.read_bytes())
    assert result["parse_status"] == "success"
    assert "sql" in result["skills"] or "python" in result["skills"]


def test_application_package(tmp_db, tmp_path):
    from docx import Document

    user = create_user("pack@example.com", "password123")
    docx_path = tmp_path / "master.docx"
    doc = Document()
    doc.add_paragraph("Profile summary for a junior data analyst.")
    doc.add_paragraph("Skills SQL Python Excel")
    doc.save(docx_path)
    store_cv_upload(user["id"], "master.docx", docx_path.read_bytes())
    job = Job(
        source="Test",
        title="Junior Data Analyst",
        company="Acme",
        location="Israel",
        published_at="2026-07-20T00:00:00+00:00",
        url="https://example.com/job",
        description="SQL Python Excel Tableau",
        score=50,
        ats_score=80,
        matched_skills="sql, python",
        missing_skills="tableau",
        match_explanation="test",
        work_type="hybrid",
        region="israel_local",
    )
    folder = create_application_package(
        user["id"],
        job,
        docx_path,
        ["sql", "python", "excel"],
    )
    assert (folder / "tailored_cv.docx").exists()
    assert (folder / "cover_letter.txt").exists()
    assert (folder / "job_description.txt").exists()
    apps = list_applications(user["id"])
    assert apps
    assert apps[0]["status"] == "Preparing"


def test_delete_account(tmp_db):
    user = create_user("del@example.com", "password123")
    delete_user_and_data(user["id"])
    assert authenticate("del@example.com", "password123") is None
